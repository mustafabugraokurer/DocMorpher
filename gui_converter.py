import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from PIL import Image, ImageSequence
from docx import Document
from pdf2docx import Converter
import pytesseract
import fitz  # PyMuPDF
import ocrmypdf
import os
import numpy as np
import subprocess
import comtypes.client
import zipfile
import xml.etree.ElementTree as ET

pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

output_dir = os.getcwd()  # Varsayilan cikti klasoru


def log(message):
    log_text.insert(tk.END, message + "\n")
    log_text.see(tk.END)

def convert_udf_to_pdf(file_path):
    try:
        base = os.path.splitext(os.path.basename(file_path))[0]
        extract_path = os.path.join(output_dir, base + "_udf")
        os.makedirs(extract_path, exist_ok=True)

        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            zip_ref.extractall(extract_path)

        xml_file = None
        for file in os.listdir(extract_path):
            if file.lower().endswith(".xml"):
                xml_file = os.path.join(extract_path, file)
                break

        if not xml_file:
            log("[HATA] XML dosyası UDF içinde bulunamadı.")
            return

        tree = ET.parse(xml_file)
        root = tree.getroot()
        text_content = "\n".join([elem.text for elem in root.iter() if elem.text])

        docx_path = os.path.join(output_dir, f"{base}.docx")
        document = Document()
        document.add_paragraph(text_content)
        document.save(docx_path)

        pdf_path = os.path.join(output_dir, f"{base}.pdf")
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()

        log(f"[OK] {pdf_path} oluşturuldu (UDF → DOCX → PDF)")
    except Exception as e:
        log(f"[HATA] {e}")

def convert_tif(file_path, output_format):
    try:
        base = os.path.splitext(os.path.basename(file_path))[0]

        if output_format == "PDF":
            img = Image.open(file_path)
            frames = [frame.convert("RGB") for frame in ImageSequence.Iterator(img)]
            output_file = os.path.join(output_dir, f"{base}.pdf")
            frames[0].save(output_file, save_all=True, append_images=frames[1:], format='PDF')

        elif output_format == "DOCX":
            doc = Document()
            doc.add_paragraph("TIF Görseli:")
            doc.add_picture(file_path)
            output_file = os.path.join(output_dir, f"{base}.docx")
            doc.save(output_file)

        elif output_format == "DOCX (OCR)":
            img = Image.open(file_path)
            doc = Document()
            for i, page in enumerate(ImageSequence.Iterator(img)):
                text = pytesseract.image_to_string(page.convert("L"), lang="tur")
                doc.add_paragraph(f"--- Sayfa {i+1} ---")
                doc.add_paragraph(text)
            output_file = os.path.join(output_dir, f"{base}_ocr.docx")
            doc.save(output_file)

        elif output_format in ["JPG", "PNG"]:
            img = Image.open(file_path).convert("RGB")
            output_file = os.path.join(output_dir, f"{base}.{output_format.lower()}")
            img.save(output_file, output_format.upper())

        else:
            raise ValueError("Bilinmeyen format secimi.")

        log(f"[OK] {output_file} olusturuldu.")
    except Exception as e:
        log(f"[HATA] {e}")


def convert_pdf_to_word(file_path):
    try:
        base = os.path.splitext(os.path.basename(file_path))[0]
        output_file = os.path.join(output_dir, f"{base}.docx")
        cv = Converter(file_path)
        cv.convert(output_file, start=0, end=None)
        cv.close()
        log(f"[OK] {output_file} olusturuldu.")
    except Exception as e:
        log(f"[HATA] {e}")


def convert_pdf_to_pdf_ocrmypdf(file_path):
    try:
        base = os.path.splitext(os.path.basename(file_path))[0]
        output_file = os.path.join(output_dir, f"{base}_ocrmypdf.pdf")
        ocrmypdf.ocr(file_path, output_file, language='tur', deskew=True)
        log(f"[OK] {output_file} olusturuldu.")
    except Exception as e:
        log(f"[HATA] {e}")


def convert_office_to_pdf(file_path):
    try:
        ext = os.path.splitext(file_path)[1].lower()
        base = os.path.splitext(os.path.basename(file_path))[0]
        output_file = os.path.join(output_dir, f"{base}.pdf")

        if ext == ".docx":
            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(file_path)
            doc.SaveAs(output_file, FileFormat=17)
            doc.Close()
            word.Quit()

        elif ext == ".xlsx":
            excel = comtypes.client.CreateObject('Excel.Application')
            wb = excel.Workbooks.Open(file_path)
            wb.ExportAsFixedFormat(0, output_file)
            wb.Close(False)
            excel.Quit()

        elif ext == ".pptx":
            powerpoint = comtypes.client.CreateObject("Powerpoint.Application")
            ppt = powerpoint.Presentations.Open(file_path, WithWindow=False)
            ppt.SaveAs(output_file, 32)  # 32 = PDF format
            ppt.Close()
            powerpoint.Quit()

        else:
            log("[UYARI] Bu Office dosya tipi desteklenmiyor.")
            return

        log(f"[OK] {output_file} olusturuldu.")
    except Exception as e:
        log(f"[HATA] {e}")





def select_file():
    filepath = filedialog.askopenfilename(filetypes=[("Desteklenen Dosyalar", "*.tif;*.tiff;*.pdf;*.udf;*.docx;*.xlsx;*.pptx")])
    if filepath:
        file_label.config(text=filepath)
        update_format_options(filepath)
        convert_button.config(state="normal")


def select_output_dir():
    global output_dir
    output_dir = filedialog.askdirectory()
    if output_dir:
        output_dir_label.config(text=f"Cikti klasoru: {output_dir}")


def update_format_options(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    menu = format_option_menu["menu"]
    menu.delete(0, "end")
    if ext in [".tif", ".tiff"]:
        valid_formats = ["PDF", "DOCX", "DOCX (OCR)", "JPG", "PNG"]
    elif ext == ".pdf":
        valid_formats = ["DOCX", "PDF (OCRmyPDF)"]
    elif ext == ".udf":
        valid_formats = ["PDF"]
    elif ext in [".docx", ".xlsx", ".pptx"]:
        valid_formats = ["PDF"]
    else:
        valid_formats = []
    for fmt in valid_formats:
        menu.add_command(label=fmt, command=lambda value=fmt: format_var.set(value))
    if valid_formats:
        format_var.set(valid_formats[0])


def convert():
    file_path = file_label.cget("text")
    selected_format = format_var.get()
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        if selected_format == "DOCX":
            convert_pdf_to_word(file_path)
        elif selected_format == "PDF (OCRmyPDF)":
            convert_pdf_to_pdf_ocrmypdf(file_path)
        else:
            log("[HATA] Bu format PDF icin desteklenmiyor.")
    elif ext in [".tif", ".tiff"]:
        convert_tif(file_path, selected_format)
    elif ext in [".docx", ".xlsx", ".pptx"]:
        convert_office_to_pdf(file_path)
    elif ext == ".udf":
        convert_udf_to_pdf(file_path)
    else:
        log("[HATA] Bu dosya tipi desteklenmiyor.")


# GUI
root = tk.Tk()
root.title("Akilli Dosya Dönüştürücü")
root.geometry("600x600")
root.configure(bg="#f4f4f4")
root.resizable(False, False)

style = ttk.Style()
style.configure("TLabel", font=("Segoe UI", 10), background="#f4f4f4")
style.configure("TButton", font=("Segoe UI", 10, "bold"), background="#0078D7", foreground="#0078D7")
style.configure("Header.TLabel", font=("Segoe UI", 12, "bold"), background="#0078D7", foreground="#002244")

# Başlık
ttk.Label(root, text="Akilli Dosya Dönüştürücü", style="Header.TLabel").pack(pady=10)

# Dosya Seçimi
ttk.Label(root, text="1. Dosyayı Seçin:").pack(pady=(5, 0))
ttk.Button(root, text="Dosya Seç", command=select_file).pack()
file_label = ttk.Label(root, text="", wraplength=500)
file_label.pack(pady=5)

# Çıktı Klasörü
ttk.Button(root, text="2. Çıktı Klasörü Seç", command=select_output_dir).pack(pady=3)
output_dir_label = ttk.Label(root, text=f"Çıktı klasörü: {output_dir}", wraplength=500)
output_dir_label.pack(pady=3)

# Bilgilendirme
ttk.Label(root, text="Desteklenen Dönüşümler:").pack(pady=(10, 0))
desc_label = ttk.Label(root, text="TIF → PDF, DOCX, DOCX (OCR), JPG, PNG\nPDF → DOCX, PDF (OCRmyPDF)\nUDF → PDF \nDOCX/XLSX/PPTX → PDF")
desc_label.pack(pady=2)

# Format Seçimi
ttk.Label(root, text="3. Hedef Formatı Seçin:").pack(pady=(10, 0))
format_var = tk.StringVar()
format_option_menu = ttk.OptionMenu(root, format_var, "PDF")
format_option_menu.pack()

# Dönüştür Butonu
convert_button = ttk.Button(root, text="Dönüştür", state="disabled", command=convert)
convert_button.pack(pady=10)

# Log Alanı
ttk.Label(root, text="İşlem Logu:").pack(pady=(10, 0))
log_text = tk.Text(root, height=10, width=70, bg="#ffffff", font=("Consolas", 9))
log_text.pack(pady=5)

root.mainloop()